import docx.document
import musicxml
from musicxml.parser.parser import _parse_node
from musicxml import *
import docx
from docx.shared import Pt
from zipfile import ZipFile
import xml.etree.ElementTree as ET
import xml
from converter import number_map, measure_map, rest_map
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import music21 as m
import math

class Measure:
    def __init__(self):
        self.key_signature = ''
        self.time_signature = ''
        self.start_bar = ''
        self.end_bar = ''
        self.notes = []

    def get_string_list(self):
        all = [self.key_signature, self.time_signature, self.start_bar] + self.notes+ [ self.end_bar]
        return [x for x in all if x != '']

WRITE_LIST = []
my_key = ''
divison = 0

def add_jianpu_paragraph(string:str, doc:docx.document.Document, text_size:int=16):
    paragraph = doc.add_paragraph(string)
    run = paragraph.runs[0]
    run.font.name = 'SimpErhuFont'
    run.font.size = Pt(text_size)
    
def unicode_to_char(unicode_string:str):
    return chr(int(unicode_string, 16))

def add_duration(note_string:str, duration:int) -> list[str]:
    """
    Given the note string and the duration, return the correct jianpu unicode character and/or added durations in a list
    """
    global divison
    frac, whole = math.modf(duration/divison)
    characters = []

    if frac > 0 and whole < 1:
        if not note_string.__contains__('-'):
            note_string += '-'
        note_string += 'l'*int(math.log(int(1/frac),2))
    characters.append(number_map[note_string])
    if frac == 0.5 and whole >= 1:
        characters.append(measure_map['dot'])
    if whole > 1:
        for x in range(1, int(whole)):
            characters.append(measure_map['dash'])

    return characters

def note_to_unicode(key:str, note:str, duration:int) -> list[str]:
    """
    Convert note text to the correct jianpu unicode characters with length
    """
    n = m.note.Note(note)
    k = m.key.Key(key)
    base_note = m.note.Note(k.getScale('major').pitches[0].name)

    interval = m.interval.Interval(base_note, n)
    distance = abs(math.floor(interval.semitones / 12)) * interval.direction

    solfage = k.getScale().getScaleDegreeFromPitch(n)

    if solfage == None:
        raise ValueError('Note not in key signature')
    if distance > 3 or distance < -3:
        raise ValueError('Note not in disatonic range')

    note_text = str(solfage)
    match distance:
        case _ if distance > 0:
            note_text += '-a'
            note_text += 'd'* abs(distance)
        case _ if distance < 0:
            note_text += '-u'
            note_text += 'd'*abs(distance)
        case _ if distance == 0:
            note_text += ''

    return add_duration(note_text, duration)

def note_to_string(note:XMLNote, duration:int) -> list[str]:
    """
    Convert XMLNote to the correct jianpu string with length
    """
    global my_key
    pitch:XMLPitch = note.get_children_of_type(XMLPitch)[0]
    step = pitch.get_children_of_type(XMLStep)[0].value_
    octave = pitch.get_children_of_type(XMLOctave)[0].value_
    duration = note.get_children_of_type(XMLDuration)[0].value_
    alter = ''
    try:
        alter = (int)(pitch.get_children_of_type(XMLAlter)[0].value_)
        match alter:
            case 1:
                alter = '#'
            case -1:
                alter = 'b'
            case 0:
                alter = ''
    except:
        pass
    
    return note_to_unicode(my_key, f'{step}{alter}{octave}', duration)

def convert_note(note:XMLNote | None)-> list[str]:
    global divison
    duration = note.get_children_of_type(XMLDuration)[0].value_
    notes = []

    # Rest
    if note.get_children_of_type(XMLRest):
        frac, whole = math.modf(duration/divison)
        rest_string = '0'
        try:
            if frac > 0 and whole < 1:
                rest_string += '-' + 'l'*int(math.log(int(1/frac),2))
            notes.append(rest_map[rest_string])
            if whole > 1:
                if frac > 0:
                    notes.append(measure_map['dot'])
                for x in range(1, int(whole)):
                    notes.append(rest_map[rest_string])
        except:
            print('Rest not found')
    # Actual note
    else:
        notes += note_to_string(note, duration)
    return notes

def convert_key_signature(key_signature:XMLKey)->str:
    """
    Converts the key signature to the correct jianpu unicode character
    """
    key:str = (str)(m.key.KeySignature(key_signature.get_children_of_type(XMLFifths)[0].value_).asKey()).split(' ')[0]
    global my_key
    my_key = key
    try:
        return measure_map['key_signatures'][key]
    except:
        print('Key signature not found, defaulting to D')
        return measure_map['key_signatures']['D']

def convert_time_signature(time_signature:XMLTime)->str:
    """
    Converts the time signature to the correct jianpu unicode character
    """
    time_signature = time_signature.get_children_of_type(XMLBeats)[0].value_+'/'+time_signature.get_children_of_type(XMLBeatType)[0].value_
    try:
        return measure_map['time_signatures'][time_signature]
    except:
        print('Time signature not found, defaulting to 4/4')
        return measure_map['time_signatures']['4/4']

def convert_attributes(attributes:XMLAttributes)->tuple[str, str]:
    """
    Returns the key and time signature of the attributes in format: (key, time)
    """
    for child in attributes.get_children():
        if isinstance(child, XMLKey):
            key = convert_key_signature(child)
        if isinstance(child, XMLTime):
            time = convert_time_signature(child)
        if isinstance(child, XMLDivisions):
            global divison
            divison = int(child.value_)
    return (key, time)

def convert_barline(barline:XMLBarline, previous:XMLMeasure | None, measure:Measure):

    if barline.get_children_of_type(XMLRepeat):

        match barline.get_children_of_type(XMLRepeat)[0].to_string().split('"')[1]:
            case 'forward':
                measure.start_bar = measure_map['bar_lines']['repeat_forward']
            case 'backward':
                measure.end_bar = measure_map['bar_lines']['repeat_backward']
            case _:
                measure.end_bar = measure_map['bar_lines']['bold_double']

        if previous is not None:
            for x in previous.get_children_of_type(XMLBarline):
                if x.get_children_of_type(XMLRepeat)[0].to_string().split('"')[1] == 'backward':
                    measure.start_bar = measure_map['bar_lines']['repeat_both']
                    break
    else:
        measure.end_bar = measure_map['bar_lines']['bold_double']

def convert_measure(measure:XMLMeasure | None):
    this_measure = Measure()
    global WRITE_LIST

    if measure is None:
        return
    for child in measure.get_children():
        if isinstance(child, XMLNote):
            # Skip erverything that is not in the first voice and staff
            if (child.get_children_of_type(XMLVoice)[0].value_ != '1'):
                continue
            if len(child.get_children_of_type(XMLStaff)) != 0 and (child.get_children_of_type(XMLStaff)[0].value_ != 1) :
                continue
            this_measure.notes += convert_note(child)
        if isinstance(child, XMLAttributes):
            this_measure.key_signature, this_measure.time_signature = convert_attributes(child)
        if isinstance(child, XMLBarline):
            convert_barline(child,measure.previous,this_measure)

    if this_measure.end_bar == '':
        this_measure.end_bar = measure_map['bar_lines']['simple']

    if this_measure.start_bar == measure_map['bar_lines']['repeat_both']:
        WRITE_LIST.pop()
    if this_measure.start_bar == measure_map['bar_lines']['repeat_forward'] and WRITE_LIST[-1] == measure_map['bar_lines']['simple']:
        WRITE_LIST.pop()
    WRITE_LIST += this_measure.get_string_list()
    convert_measure(measure.next)

# Get xml file
with ZipFile('mxl/sample.mxl', 'r') as zipObj:
    xml_string = zipObj.read('score.xml').decode('utf-8')

xml = ET.fromstring(xml_string)

mxl:XMLScorePartwise  = _parse_node(xml)

credit_list: list = {}
for credit in mxl.get_children_of_type(XMLCredit):
    type = value = ''
    for child in credit.get_children():
        if isinstance(child, XMLCreditType):
            type = child.value_
        if isinstance(child, XMLCreditWords):
            value = child.value_
    credit_list[type] = value

parts:list[XMLPart] = mxl.get_children_of_type(XMLPart)


# # Create a new Document
path = "docx/sample.docx"
doc = docx.Document(path)


convert_measure(parts[0].get_children()[0])

# Replace the title and composer
doc.paragraphs[0].text = doc.paragraphs[0].text.replace('Song Title', credit_list['title'])
doc.paragraphs[1].text = doc.paragraphs[1].text.replace('Composer', credit_list['composer'])

new_string = ''.join(unicode_to_char(char) for char in WRITE_LIST)

add_jianpu_paragraph(new_string, doc)
doc.save("docx/new.docx")